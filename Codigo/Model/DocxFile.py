import textract
import docx
from abc import ABC
from Codigo.Model.File import File


def extract_text_from_table(table):
    table_text = ""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                table_text += paragraph.text + "\n"
    return table_text


class DocxFile(File, ABC):
    def __init__(self, name, url_file):
        super().__init__(name, url_file)

    def get_text(self):
        try:
            text = textract.process(self.url_file, method='docx', encoding='utf-8').decode('utf-8')
            text = self.verify_footer(text)

            # Escribe el texto en un archivo de texto
            path = self.path + self.name + ".txt"
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
            return {
                'response': False,
                'message': text,

            }
        except FileNotFoundError as e:
            return {
                'response': True,
                'message': f"Error: El archivo {self.url_file} no se encontró.",
            }
        except Exception as e:
            return {
                'response': True,
                'message': "Error: No se logró cargar el archivo.",
            }

    def verify_footer(self, text):
        last_paragraph = ""
        last_table = ""

        # Abre el archivo DOCX
        doc = docx.Document(self.url_file)

        # Inicializa un índice para recorrer los párrafos y las tablas hacia atrás
        index = -1

        if len(doc.paragraphs) > 0:
            # Bucle while para recorrer los párrafos hacia atrás
            while index >= -len(doc.paragraphs):
                para = doc.paragraphs[index]
                # Comprueba si el párrafo es parte del pie de página o si está vacío
                if para.text.strip() != "" and para.text != "\x0c":
                    # Agrega el texto del párrafo al principio de la cadena (hacia atrás)
                    last_paragraph = para.text
                    break
                index -= 1

            last_paragraph_word = last_paragraph.split()[-1]
            index1 = text.index(last_paragraph_word)
            # print(last_paragraph.split()[-1])
            # print("xxxxxxxxxxxxxxxxxxxxxxxxxxxxxx")
        else:
            index1 = -1

        # Inicializa un índice para recorrer las tablas hacia atrás
        index = -1

        # Bucle while para recorrer las tablas hacia atrás
        if len(doc.tables) > 0:
            while index >= -len(doc.tables):
                table = doc.tables[index]
                t = extract_text_from_table(table)

                if t.strip() != "" and t != "\x0c":
                    # Agrega el texto de la tabla al principio de la cadena (hacia atrás)
                    last_table = t
                    break
                index -= 1

            last_table_word = last_table.split()[-1]
            index2 = text.index(last_table_word)
        else:
            index2 = -1

        if index1 > index2:
            text = text.split(last_paragraph_word)[0] + last_paragraph_word
        elif index2 > index1:
            text = text.split(last_table_word)[0] + last_table_word

        return text


x = DocxFile("2", "2.docx")
print(x.get_text())
